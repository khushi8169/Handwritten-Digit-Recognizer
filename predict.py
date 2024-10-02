import numpy as np
import tensorflow as tf
import cv2
from docx import Document
from docx.shared import Inches


def predict():


    model = tf.keras.models.load_model('digits_model.keras')


    img = cv2.imread('captured_image.png', cv2.IMREAD_GRAYSCALE)
    if img is None:
        print("Image not found or cannot be read.")
        return


    _, img = cv2.threshold(img, 128, 255, cv2.THRESH_BINARY_INV)
    img = tf.keras.utils.normalize(img, axis=1)
    img = np.expand_dims(img, axis=-1)
    img = np.expand_dims(img, axis=0)


    prediction = model.predict(img)
    print(f"Predicted digit: {np.argmax(prediction)}")


    doc = Document()
    doc.add_heading('Digit Recognition Result')
    doc.add_paragraph(f'Predicted Digit: {np.argmax(prediction)}')


    doc.save('prediction_result.docx')


if __name__ == "__main__":
    predict()
